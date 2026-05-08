"""MIME detection and text extraction utilities for local file parsing."""

from __future__ import annotations

import logging
from collections.abc import Generator
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlsplit

logger = logging.getLogger(__name__)


@dataclass
class ParsedFile:
    """Result of parsing a local file."""

    mime_type: str
    text_content: str
    is_binary: bool
    file_size_bytes: int = 0
    encoding: str | None = None
    parse_error: str | None = None


@dataclass
class ParsedBytes:
    """Result of parsing in-memory bytes."""

    mime_type: str
    raw_content: str
    text_content: str
    is_binary: bool
    file_size_bytes: int
    parse_error: str | None = None


_TEXT_RAW_MIME_TYPES = {
    "application/json",
    "application/xml",
    "text/xml",
    "application/xhtml+xml",
}

_TABULAR_MIME_TYPES = {
    "text/csv",
    "text/tab-separated-values",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/parquet",
    "application/vnd.apache.parquet",
}

_MIME_HINTS_BY_EXTENSION = {
    ".csv": "text/csv",
    ".tsv": "text/tab-separated-values",
    ".parquet": "application/parquet",
    ".json": "application/json",
    ".xml": "application/xml",
    ".html": "text/html",
    ".htm": "text/html",
    ".md": "text/markdown",
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def _normalize_mime_type(mime_type: str | None) -> str:
    if not mime_type:
        return ""
    return str(mime_type).split(";", 1)[0].strip().lower()


def _file_extension(file_name: str) -> str:
    if not file_name:
        return ""
    path = urlsplit(file_name).path
    value = path if path else file_name
    return Path(value).suffix.lower()


def infer_mime_type_from_file_name(file_name: str) -> str:
    """Infer MIME type from file name or URL path extension."""
    extension = _file_extension(file_name)
    return _MIME_HINTS_BY_EXTENSION.get(extension, "application/octet-stream")


def normalize_detected_mime_type(detected_mime_type: str, file_name: str) -> str:
    """
    Normalize detected MIME with filename-based fallbacks.

    Keeps parser behavior stable for sources that declare generic or plain-text
    content-types for tabular files.
    """
    mime = _normalize_mime_type(detected_mime_type)
    inferred_mime = infer_mime_type_from_file_name(file_name)

    if not mime or mime == "application/octet-stream":
        return inferred_mime

    if mime == "text/plain" and inferred_mime in _TABULAR_MIME_TYPES:
        return inferred_mime

    return mime


def _is_text_like_mime_type(mime_type: str) -> bool:
    normalized_mime = _normalize_mime_type(mime_type)
    return normalized_mime.startswith("text/") or normalized_mime in _TEXT_RAW_MIME_TYPES


def _detect_magic_mime_type(file_bytes: bytes) -> str | None:
    signatures: tuple[tuple[bytes, str], ...] = (
        (b"\x89PNG\r\n\x1a\n", "image/png"),
        (b"%PDF-", "application/pdf"),
        (b"\xff\xd8\xff", "image/jpeg"),
        (b"GIF87a", "image/gif"),
        (b"GIF89a", "image/gif"),
        (b"PK\x03\x04", "application/zip"),
    )

    for signature, mime_type in signatures:
        if file_bytes.startswith(signature):
            return mime_type

    return None


def _sniff_text_mime(file_bytes: bytes) -> str:
    """Fallback MIME detection for text formats not handled by filetype."""
    # Check for null bytes → binary
    if b"\x00" in file_bytes[:8192]:
        return "application/octet-stream"

    # Try to decode a sample for text-based sniffing
    sample = ""
    try:
        import chardet

        detected = chardet.detect(file_bytes[:4096])
        encoding = detected.get("encoding") or "utf-8"
        sample = file_bytes[:4096].decode(encoding, errors="replace")
    except Exception:
        try:
            sample = file_bytes[:4096].decode("utf-8", errors="replace")
        except Exception:
            return "application/octet-stream"

    stripped = sample.lstrip()

    if stripped.startswith("{") or stripped.startswith("["):
        return "application/json"
    if stripped.startswith("<?xml"):
        return "application/xml"
    if stripped.lower().startswith("<!doctype html") or stripped.lower().startswith("<html"):
        return "text/html"

    # CSV heuristic: first non-empty line has multiple commas
    first_line = stripped.split("\n")[0] if "\n" in stripped else stripped
    if first_line.count(",") >= 2:
        return "text/csv"

    return "text/plain"


def detect_mime_type(file_bytes: bytes) -> str:
    """
    Detect MIME type from file bytes.

    Uses magic-byte detection first (filetype library), then falls back to
    text-based sniffing for formats that filetype doesn't cover.
    """
    if not file_bytes:
        return "application/octet-stream"

    magic_mime_type = _detect_magic_mime_type(file_bytes)
    if magic_mime_type:
        return magic_mime_type

    try:
        import filetype

        kind = filetype.guess(file_bytes)
        if kind is not None:
            return str(kind.mime)
    except Exception as e:
        logger.debug(f"filetype detection failed: {e}")

    return _sniff_text_mime(file_bytes)


def extract_text(file_bytes: bytes, mime_type: str) -> tuple[str, str | None]:
    """
    Extract plain text from file bytes based on MIME type.

    Returns:
        (text_content, error_message_or_None)
    """
    # Binary media types — no text extraction
    if mime_type.startswith(("image/", "audio/", "video/")):
        return "", None

    # PDF
    if mime_type == "application/pdf":
        try:
            import io

            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text:
                        pages.append(text)
            return "\n\n".join(pages), None
        except Exception as e:
            return "", f"PDF extraction failed: {e}"

    # DOCX
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import io

            import docx

            doc = docx.Document(io.BytesIO(file_bytes))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts), None
        except Exception as e:
            return "", f"DOCX extraction failed: {e}"

    # XLSX
    if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        try:
            import io

            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            rows: list[str] = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows.append("\t".join(cells))
            return "\n".join(rows), None
        except Exception as e:
            return "", f"XLSX extraction failed: {e}"

    # HTML / XHTML
    if mime_type in ("text/html", "application/xhtml+xml"):
        try:
            from .content_extraction import html_to_text

            text = _decode_bytes(file_bytes)
            return html_to_text(text), None
        except Exception as e:
            return "", f"HTML extraction failed: {e}"

    # JSON, XML — decode as-is
    if mime_type in (
        "application/json",
        "application/xml",
        "text/xml",
    ):
        return _decode_bytes(file_bytes), None

    # Plain text, CSV, Markdown, and other text/* types
    if mime_type.startswith("text/"):
        return _decode_bytes(file_bytes), None

    # Parquet
    if mime_type in ("application/parquet", "application/vnd.apache.parquet"):
        try:
            import io

            import pyarrow.parquet as pq  # type: ignore[import-not-found, import-untyped]

            table = pq.read_table(io.BytesIO(file_bytes))
            column_names = table.schema.names
            lines: list[str] = []
            for row_index in range(table.num_rows):
                lines.append(f"row_{row_index + 1}:")
                for col in column_names:
                    col_array = table.column(col)
                    cell = col_array[row_index].as_py()
                    cell_str = "" if cell is None else str(cell)
                    rendered_lines = cell_str.splitlines() or [""]
                    first_line, *continuation_lines = rendered_lines
                    lines.append(f"  {col}: {first_line}")
                    for cont in continuation_lines:
                        lines.append(f"    {cont}")
                lines.append("")
            return "\n".join(lines), None
        except Exception as e:
            return "", f"Parquet extraction failed: {e}"

    # Unknown / binary
    return "", None


def _decode_bytes(file_bytes: bytes) -> str:
    """Decode bytes to str using chardet for encoding detection."""
    try:
        import chardet

        detected = chardet.detect(file_bytes[:65536])
        encoding = detected.get("encoding") or "utf-8"
        return file_bytes.decode(encoding, errors="replace")
    except Exception:
        return file_bytes.decode("utf-8", errors="replace")


def resolve_mime_type(
    file_bytes: bytes,
    *,
    declared_mime_type: str | None = None,
    file_name: str = "",
) -> str:
    """
    Resolve effective MIME type from declared hint, magic-byte detection, and file extension.

    Kept separate from full parsing so callers can detect format cheaply without
    paying for text extraction (e.g. when content will be streamed in pages later).
    """
    declared_mime = _normalize_mime_type(declared_mime_type)
    detected_mime = _normalize_mime_type(detect_mime_type(file_bytes))
    inferred_mime = infer_mime_type_from_file_name(file_name)

    if declared_mime and declared_mime != "application/octet-stream":
        mime_type = declared_mime
    elif detected_mime and detected_mime != "application/octet-stream":
        mime_type = detected_mime
    elif inferred_mime and inferred_mime != "application/octet-stream":
        mime_type = inferred_mime
    else:
        mime_type = declared_mime or detected_mime or inferred_mime or "application/octet-stream"

    mime_type = normalize_detected_mime_type(mime_type, file_name)
    if mime_type == "application/octet-stream" and inferred_mime != "application/octet-stream":
        mime_type = inferred_mime

    return mime_type


def parse_bytes(
    file_bytes: bytes,
    *,
    declared_mime_type: str | None = None,
    file_name: str = "",
) -> ParsedBytes:
    """
    Parse in-memory bytes: resolve MIME type and extract raw/text content.

    Used by the sandbox and any caller that needs a complete ParsedBytes in one shot.
    Object-storage sources prefer resolve_mime_type() + iter_file_pages() to avoid
    loading all content into memory before detector scanning.
    """
    file_size_bytes = len(file_bytes)
    mime_type = resolve_mime_type(
        file_bytes, declared_mime_type=declared_mime_type, file_name=file_name
    )

    text_content, parse_error = extract_text(file_bytes, mime_type)
    raw_content = _decode_bytes(file_bytes) if _is_text_like_mime_type(mime_type) else ""

    if mime_type in {"text/html", "application/xhtml+xml"} and raw_content and not text_content:
        from .content_extraction import html_to_text

        text_content = html_to_text(raw_content)

    is_binary = (
        mime_type.startswith(("image/", "audio/", "video/"))
        or mime_type == "application/octet-stream"
    )

    return ParsedBytes(
        mime_type=mime_type,
        raw_content=raw_content,
        text_content=text_content,
        is_binary=is_binary,
        file_size_bytes=file_size_bytes,
        parse_error=parse_error,
    )


def iter_file_pages(
    file_bytes: bytes,
    mime_type: str,
    batch_size: int = 100,
    include_column_names: bool = True,
) -> Generator[str, None, None]:
    """
    Iterate over file content in pages of up to batch_size rows or lines.

    Parquet / CSV / TSV  → yields batch_size *rows* per page with labelled columns.
    All other extractable types (PDF, DOCX, TXT, JSON, XML, XLSX, …) → extracts the
    full text once via extract_text(), then yields batch_size *lines* per page.
    Non-extractable types (images, audio, video, unknown binary) → yields nothing.

    New file formats only need to be added to extract_text() — not here.
    """
    normalized = _normalize_mime_type(mime_type)

    if normalized in ("application/parquet", "application/vnd.apache.parquet"):
        yield from _iter_parquet_pages(file_bytes, batch_size, include_column_names)
    elif normalized in ("text/csv", "text/tab-separated-values"):
        yield from _iter_csv_pages(file_bytes, batch_size, include_column_names)
    else:
        text, error = extract_text(file_bytes, normalized)
        if error:
            logger.warning("Text extraction error (%s): %s", mime_type, error)
        if text:
            yield from _iter_text_lines(text, batch_size)


def _iter_text_lines(text: str, batch_size: int) -> Generator[str, None, None]:
    """Yield non-empty text in chunks of batch_size lines."""
    lines = text.splitlines(keepends=True)
    for start in range(0, len(lines), batch_size):
        chunk = "".join(lines[start : start + batch_size])
        if chunk.strip():
            yield chunk


_PARQUET_MAGIC = b"PAR1"


def _iter_parquet_pages(
    file_bytes: bytes,
    batch_size: int,
    include_column_names: bool,
) -> Generator[str, None, None]:
    # Parquet files begin AND end with the 4-byte magic "PAR1".  If the footer
    # is missing the bytes were truncated mid-download; pyarrow's C++ thread
    # pool will hang indefinitely trying to read schema metadata that isn't
    # there, locking all worker threads on a futex.  Bail out early instead.
    if len(file_bytes) < 8 or file_bytes[-4:] != _PARQUET_MAGIC:
        logger.warning(
            "Parquet bytes appear truncated (footer magic missing, %d bytes); skipping",
            len(file_bytes),
        )
        return

    try:
        import io

        import pyarrow.parquet as pq  # type: ignore[import-not-found, import-untyped]

        # ParquetFile + iter_batches() reads one row-group at a time instead of
        # loading the whole table into memory, and surfaces schema errors early
        # (before reading any data) so a bad file can't lock the C++ thread pool.
        pf = pq.ParquetFile(io.BytesIO(file_bytes))
        abs_row = 0
        for batch in pf.iter_batches(batch_size=batch_size):
            col_names = batch.schema.names
            lines: list[str] = []
            for local_idx in range(batch.num_rows):
                lines.append(f"row_{abs_row + 1}:")
                for col_i, col in enumerate(col_names):
                    cell = batch.column(col_i)[local_idx].as_py()
                    cell_str = "" if cell is None else str(cell)
                    first, *rest = cell_str.splitlines() or [""]
                    lines.append(f"  {col}: {first}" if include_column_names else f"  {first}")
                    lines.extend(f"    {c}" for c in rest)
                lines.append("")
                abs_row += 1
            if lines:
                yield "\n".join(lines)
    except Exception as exc:
        logger.warning("Parquet page iteration failed: %s", exc)


def _iter_csv_pages(
    file_bytes: bytes,
    batch_size: int,
    include_column_names: bool,
) -> Generator[str, None, None]:
    import csv
    import io

    try:
        text = _decode_bytes(file_bytes)
        reader = csv.DictReader(io.StringIO(text))
        headers = list(reader.fieldnames or [])

        batch: list[dict[str, str]] = []
        total_seen = 0

        for row in reader:
            batch.append(dict(row))
            total_seen += 1
            if len(batch) >= batch_size:
                yield _format_tabular_page(
                    batch, headers, total_seen - len(batch) + 1, include_column_names
                )
                batch = []

        if batch:
            yield _format_tabular_page(
                batch, headers, total_seen - len(batch) + 1, include_column_names
            )
    except Exception as exc:
        logger.warning("CSV page iteration failed: %s", exc)


def _format_tabular_page(
    rows: list[dict[str, str]],
    headers: list[str],
    abs_row_start: int,
    include_column_names: bool,
) -> str:
    lines: list[str] = []
    for i, row in enumerate(rows):
        lines.append(f"row_{abs_row_start + i}:")
        for col in headers:
            first, *rest = (row.get(col) or "").splitlines() or [""]
            lines.append(f"  {col}: {first}" if include_column_names else f"  {first}")
            lines.extend(f"    {c}" for c in rest)
        lines.append("")
    return "\n".join(lines)


def parse_file(file_path: Path) -> ParsedFile:
    """
    Parse a local file: detect MIME type and extract text.

    Args:
        file_path: Path to the file on disk.

    Returns:
        ParsedFile with mime_type, text_content, is_binary, etc.

    Raises:
        FileNotFoundError: If file_path does not exist.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_bytes = file_path.read_bytes()
    parsed = parse_bytes(file_bytes, file_name=file_path.name)

    encoding: str | None = None
    if not parsed.is_binary and parsed.text_content:
        try:
            import chardet

            detected = chardet.detect(file_bytes[:65536])
            encoding = detected.get("encoding")
        except Exception:
            pass

    return ParsedFile(
        mime_type=parsed.mime_type,
        text_content=parsed.text_content,
        is_binary=parsed.is_binary,
        file_size_bytes=parsed.file_size_bytes,
        encoding=encoding,
        parse_error=parsed.parse_error,
    )
