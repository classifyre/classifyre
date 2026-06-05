"""Normalized file metadata extraction.

Produces a dict of normalized metadata keys (``size_bytes``, ``row_count``,
``page_count`` ...) that mean the same thing regardless of which source the
file came from (object storage, Confluence/Jira attachment, Notion file, ...).

All extraction is best-effort: any failure is captured under ``parse_error``
and a partial dict is returned. This function never raises.

The optional parsing libraries (pdfplumber, openpyxl, pyarrow, chardet, PIL)
live in the ``file-processing`` / ``content`` uv groups and are imported
lazily so this module is importable even when they are absent.
"""

from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit

logger = logging.getLogger(__name__)

# Cap byte scanning so a huge file does not force a full read just for metadata.
_CSV_MAX_SCAN_BYTES = 5 * 1024 * 1024

# Per-content-type key sets. These line up 1:1 with the reusable contentTypes in
# the x-asset-metadata catalog; the conformance test asserts the union equals
# the catalog's file content types.
FILE_BASE_KEYS: frozenset[str] = frozenset({"size_bytes", "mime_type", "parse_error"})
IMAGE_KEYS: frozenset[str] = frozenset({"image_width", "image_height"})
DOCUMENT_KEYS: frozenset[str] = frozenset({"page_count", "paragraph_count", "table_count"})
SPREADSHEET_KEYS: frozenset[str] = frozenset({"row_count", "columns", "encoding"})
TEXT_KEYS: frozenset[str] = frozenset({"encoding"})
JSON_KEYS: frozenset[str] = frozenset({"json_root_type", "top_level_keys", "array_length"})

# Every key extract_file_metadata may emit (union of all content-type key sets).
FILE_METADATA_KEYS: frozenset[str] = (
    FILE_BASE_KEYS | IMAGE_KEYS | DOCUMENT_KEYS | SPREADSHEET_KEYS | TEXT_KEYS | JSON_KEYS
)


def _normalize_mime(mime_type: str) -> str:
    return (mime_type or "").split(";", 1)[0].strip().lower()


def _extension(file_name: str) -> str:
    if not file_name:
        return ""
    path = urlsplit(file_name).path or file_name
    return Path(path).suffix.lower()


def extract_file_metadata(
    file_bytes: bytes,
    mime_type: str,
    *,
    file_name: str = "",
) -> dict[str, Any]:
    """Return normalized metadata for a file's raw bytes.

    Always sets ``size_bytes`` and ``mime_type``. Adds format-specific keys
    (``page_count``, ``row_count``, ``column_count``, ``column_names``,
    ``encoding``, ``image_width``, ``image_height``) where they can be derived.
    On any extraction error, sets ``parse_error`` and returns what was gathered.
    """
    metadata: dict[str, Any] = {
        "size_bytes": len(file_bytes),
    }
    normalized = _normalize_mime(mime_type)
    if normalized:
        metadata["mime_type"] = normalized
    extension = _extension(file_name)

    try:
        if normalized == "application/pdf" or extension == ".pdf":
            metadata.update(_pdf_metadata(file_bytes))
        elif _is_docx(normalized, extension):
            metadata.update(_docx_metadata(file_bytes))
        elif _is_parquet(normalized, extension):
            metadata.update(_parquet_metadata(file_bytes))
        elif _is_xlsx(normalized, extension):
            metadata.update(_xlsx_metadata(file_bytes))
        elif _is_delimited(normalized, extension):
            metadata.update(_csv_metadata(file_bytes, extension))
        elif normalized.startswith("image/") or extension in _IMAGE_EXTENSIONS:
            metadata.update(_image_metadata(file_bytes))
        elif _is_json(normalized, extension):
            metadata.update(_json_metadata(file_bytes))
        elif normalized.startswith("text/") or extension in _TEXT_EXTENSIONS:
            metadata.update(_encoding_metadata(file_bytes))
    except Exception as exc:  # never raise - metadata is best-effort
        logger.debug("File metadata extraction failed for %s: %s", file_name, exc)
        metadata["parse_error"] = str(exc)

    return metadata


_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".ico"}
_TEXT_EXTENSIONS = {".txt", ".md", ".xml", ".log", ".yaml", ".yml"}


def _is_parquet(mime: str, extension: str) -> bool:
    return "parquet" in mime or extension == ".parquet"


def _is_xlsx(mime: str, extension: str) -> bool:
    return (
        mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        or extension == ".xlsx"
    )


def _is_docx(mime: str, extension: str) -> bool:
    return (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or extension == ".docx"
    )


def _is_delimited(mime: str, extension: str) -> bool:
    return mime in ("text/csv", "text/tab-separated-values") or extension in (".csv", ".tsv")


def _is_json(mime: str, extension: str) -> bool:
    return mime in ("application/json", "text/json") or extension == ".json"


def _docx_metadata(file_bytes: bytes) -> dict[str, Any]:
    import docx  # type: ignore[import-untyped]

    document = docx.Document(io.BytesIO(file_bytes))
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
    }


def _json_metadata(file_bytes: bytes) -> dict[str, Any]:
    text = file_bytes[:_CSV_MAX_SCAN_BYTES].decode("utf-8", errors="replace")
    data = json.loads(text)
    if isinstance(data, dict):
        return {"json_root_type": "object", "top_level_keys": len(data)}
    if isinstance(data, list):
        return {"json_root_type": "array", "array_length": len(data)}
    return {"json_root_type": "scalar"}


def _pdf_metadata(file_bytes: bytes) -> dict[str, Any]:
    import pdfplumber  # type: ignore[import-untyped]

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        return {"page_count": len(pdf.pages)}


def build_columns(names: list[str], types: dict[str, str] | None = None) -> list[dict[str, str]]:
    """Build the normalized ``columns`` list of ``{name, type}`` objects.

    ``type`` is "" when the format does not expose column types (csv/xlsx).
    """
    type_map = types or {}
    return [{"name": name, "type": type_map.get(name, "")} for name in names]


def _parquet_metadata(file_bytes: bytes) -> dict[str, Any]:
    import pyarrow.parquet as pq  # type: ignore[import-untyped]

    parquet_file = pq.ParquetFile(io.BytesIO(file_bytes))
    schema = parquet_file.schema_arrow
    column_names = list(schema.names)
    column_types = {field.name: str(field.type) for field in schema}
    return {
        "row_count": parquet_file.metadata.num_rows,
        "columns": build_columns(column_names, column_types),
    }


def _xlsx_metadata(file_bytes: bytes) -> dict[str, Any]:
    import openpyxl  # type: ignore[import-untyped]

    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    try:
        sheet = workbook.worksheets[0]
        header_row = next(sheet.iter_rows(values_only=True), None)
        column_names = [str(cell) for cell in header_row if cell is not None] if header_row else []
        # max_row/max_column are reliable on read-only sheets after dimension scan.
        row_count = (sheet.max_row or 1) - 1  # exclude header row
        result: dict[str, Any] = {"row_count": max(row_count, 0)}
        if column_names:
            result["columns"] = build_columns(column_names)
        return result
    finally:
        workbook.close()


def _csv_metadata(file_bytes: bytes, extension: str) -> dict[str, Any]:
    encoding_meta = _encoding_metadata(file_bytes)
    encoding = encoding_meta.get("encoding") or "utf-8"
    delimiter = "\t" if extension == ".tsv" else ","
    # Wrap bytes in a TextIOWrapper so csv.reader decodes lazily line-by-line,
    # counting every row in the file without building one large decoded string.
    # This also handles multi-line quoted fields correctly across chunk boundaries.
    stream = io.TextIOWrapper(io.BytesIO(file_bytes), encoding=encoding, errors="replace")
    reader = csv.reader(stream, delimiter=delimiter)
    column_names: list[str] = []
    row_count = 0
    for index, row in enumerate(reader):
        if index == 0:
            column_names = [cell.strip() for cell in row]
            continue
        row_count += 1
    result: dict[str, Any] = {"row_count": row_count}
    if column_names:
        result["columns"] = build_columns(column_names)
    result.update(encoding_meta)
    return result


def _image_metadata(file_bytes: bytes) -> dict[str, Any]:
    from PIL import Image  # type: ignore[import-untyped]

    with Image.open(io.BytesIO(file_bytes)) as image:
        return {"image_width": image.width, "image_height": image.height}


def _encoding_metadata(file_bytes: bytes) -> dict[str, Any]:
    try:
        import chardet  # type: ignore[import-untyped]
    except ImportError:
        return {}
    detected = chardet.detect(file_bytes[:65536])
    encoding = detected.get("encoding") if isinstance(detected, dict) else None
    return {"encoding": encoding} if encoding else {}
